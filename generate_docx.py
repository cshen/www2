#!/usr/bin/env -S uv run --script
# /// script
# requires-python = ">=3.9"
# dependencies = [
#   "python-docx",
#   "bibtexparser",
# ]
# ///
"""
generate_docx.py — Formal resume-style publication list Word document.

Reads cs.bib → writes publications.docx.
Same style as generate_pdf.py: Candara font, numbered entries,
blue titles, italic venue, C. Shen underlined.

Usage:
    python3 generate_docx.py [output.docx]
    uv run generate_docx.py [output.docx]
"""

import re
import sys
from collections import OrderedDict
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import AUTHOR_NAME
from generate import parse_bib, apply_title_case, pub_venue_str, note_clean

OUTPUT_DOCX = sys.argv[1] if len(sys.argv) > 1 else "publications.docx"

# ---------------------------------------------------------------------------
# Colours
# ---------------------------------------------------------------------------

C_INK   = RGBColor(0x1c, 0x1c, 0x1c)
C_LIGHT = RGBColor(0x88, 0x88, 0x88)
C_TITLE = RGBColor(0x1a, 0x56, 0xdb)
C_NUM   = RGBColor(0x1a, 0x56, 0xdb)
C_NOTE  = RGBColor(0x1a, 0x56, 0xdb)
C_HDR   = RGBColor(0x00, 0x30, 0x87)

F_BODY  = "Candara"
FS      = 9.5      # body font size (pt)
NUM_W   = 20       # hanging indent width (pt)
LEAD    = 13.0     # exact line spacing (pt)

# ---------------------------------------------------------------------------
# Shared helpers (duplicated from generate_pdf.py to keep script standalone)
# ---------------------------------------------------------------------------

_PARTICLES = frozenset({
    'van','de','den','der','von','di','del','le','la','los','las',
    'ter','ten','af','da','dos','du','bin','binte',
})
MAX_AUTHORS = 999

def _abbrev_one(raw: str) -> str:
    raw = re.sub(r'[{}]', '', raw).strip()
    raw = re.sub(r'\s+', ' ', raw)
    if not raw: return ''
    if ',' in raw:
        idx    = raw.index(',')
        last   = raw[:idx].strip()
        firsts = raw[idx+1:].strip()
        initials = ''.join(w[0].upper() + '.' for w in firsts.split() if w and w[0].isalpha())
        return f"{initials} {last}" if initials else last
    tokens = raw.split()
    if len(tokens) == 1: return tokens[0]
    if len(tokens) == 2: return f"{tokens[0][0].upper()}. {tokens[1]}"
    last_start = len(tokens) - 1
    for i in range(1, len(tokens)):
        if tokens[i][0].islower() or tokens[i].lower() in _PARTICLES:
            last_start = i
            break
    first_tokens = tokens[:last_start]
    last_tokens  = tokens[last_start:]
    initials = ''.join(t[0].upper() + '.' for t in first_tokens if t and t[0].isalpha())
    return f"{initials} {' '.join(last_tokens)}" if initials else ' '.join(last_tokens)

def _is_primary(name: str) -> bool:
    return name.strip() == AUTHOR_NAME

_GREEK = {
    'alpha':'α','beta':'β','gamma':'γ','delta':'δ','epsilon':'ε',
    'zeta':'ζ','eta':'η','theta':'θ','iota':'ι','kappa':'κ',
    'lambda':'λ','mu':'μ','nu':'ν','xi':'ξ','pi':'π','rho':'ρ',
    'sigma':'σ','tau':'τ','upsilon':'υ','phi':'φ','chi':'χ',
    'psi':'ψ','omega':'ω','Alpha':'Α','Beta':'Β','Gamma':'Γ',
    'Delta':'Δ','Theta':'Θ','Lambda':'Λ','Xi':'Ξ','Pi':'Π',
    'Sigma':'Σ','Phi':'Φ','Psi':'Ψ','Omega':'Ω',
}
_SUPER = {'0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵',
          '6':'⁶','7':'⁷','8':'⁸','9':'⁹','+':'⁺','-':'⁻','n':'ⁿ'}
_SUB   = {'0':'₀','1':'₁','2':'₂','3':'₃','4':'₄','5':'₅',
          '6':'₆','7':'₇','8':'₈','9':'₉','n':'ₙ','i':'ᵢ','k':'ₖ'}
_SYMS  = {
    r'\infty':'∞',r'\times':'×',r'\cdot':'·',r'\ldots':'…',r'\dots':'…',
    r'\nabla':'∇',r'\partial':'∂',r'\leq':'≤',r'\geq':'≥',r'\neq':'≠',
    r'\approx':'≈',r'\rightarrow':'→',r'\leftarrow':'←',r'\to':'→',
    r'\in':'∈',r'\subset':'⊂',r'\cup':'∪',r'\cap':'∩',r'\pm':'±',r'\sqrt':'√',
}

def _math_to_unicode(s: str) -> str:
    for name, sym in _GREEK.items():
        s = re.sub(r'\\' + name + r'(?![a-zA-Z])', sym, s)
    for cmd, sym in _SYMS.items():
        s = s.replace(cmd, sym)
    def _sup(m): return ''.join(_SUPER.get(c, c) for c in (m.group(1) or m.group(2)))
    s = re.sub(r'\^\{([^}]*)\}|\^([^\s{\\])', _sup, s)
    def _sub(m): return ''.join(_SUB.get(c, c) for c in (m.group(1) or m.group(2)))
    s = re.sub(r'_\{([^}]*)\}|_([^\s{\\])', _sub, s)
    s = re.sub(r'\\[a-zA-Z]+', '', s)
    return re.sub(r'[{}]', '', s).strip()

def _clean_title(raw: str) -> str:
    if not raw: return ''
    parts = re.split(r'(\$[^$]+\$)', raw)
    out = []
    for p in parts:
        if p.startswith('$') and p.endswith('$') and len(p) > 2:
            out.append(_math_to_unicode(p[1:-1]))
        else:
            p = re.sub(r'\{([^{}]*)\}', r'\1', p)
            p = p.replace('{', '').replace('}', '')
            out.append(p)
    return ''.join(out).strip()

def _prep_title(raw: str) -> str:
    """Normalise whitespace and collapse single-letter brace groups (e.g. {B}ezier → Bezier)
    BEFORE apply_title_case runs, so title casing works on whole words."""
    raw = re.sub(r'\s+', ' ', raw).strip()
    raw = re.sub(r'\{([A-Za-z])\}', r'\1', raw)
    return raw

_ABBREV_MAP = [
    (re.compile(r'\bTransactions on\b', re.I), 'Trans.'),
    (re.compile(r'\bConference on\b',   re.I), 'Conf.'),
    (re.compile(r'\bJournal of\b',      re.I), 'J.'),
    (re.compile(r'\bInternational\b',   re.I), "Int'l"),
    (re.compile(r'\b and \b',           re.I), ' & '),
]

def _shorten_venue(s: str) -> str:
    for pat, repl in _ABBREV_MAP:
        s = pat.sub(repl, s)
    return s

_CONF_TYPES = {'inproceedings', 'inproceedgins', 'conference'}
_JOUR_TYPES = {'article'}

def _entry_kind(e: dict) -> int:
    t = (e.get('ENTRYTYPE') or '').lower().strip()
    if t in _CONF_TYPES: return 0
    if t in _JOUR_TYPES: return 1
    return 2

# ---------------------------------------------------------------------------
# Run helper
# ---------------------------------------------------------------------------

def _run(para, text: str, size=None, color=None, italic=False, underline=False):
    if not text: return
    run = para.add_run(text)
    run.font.name  = F_BODY
    run.font.size  = Pt(size or FS)
    run.font.color.rgb = color or C_INK
    run.font.italic    = italic
    run.font.underline = underline
    run.font.bold      = False
    return run

# ---------------------------------------------------------------------------
# Entry paragraph
# ---------------------------------------------------------------------------

def add_entry(doc: Document, e: dict, n: int):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.left_indent        = Pt(NUM_W)
    pf.first_line_indent  = Pt(-NUM_W)
    pf.space_after        = Pt(2)
    pf.space_before       = Pt(0)
    pf.line_spacing_rule  = WD_LINE_SPACING.EXACTLY
    pf.line_spacing       = Pt(LEAD)

    # Index number (blue, smaller)
    _run(para, f"{n}.", size=8, color=C_NUM)
    _run(para, "  ")

    # Authors
    author_str = e.get('author') or ''
    if author_str:
        raw_names = re.split(r'\s+and\s+', author_str.strip(), flags=re.IGNORECASE)
        for i, raw in enumerate(raw_names[:MAX_AUTHORS]):
            raw  = re.sub(r'\s+', ' ', raw).strip()
            abbr = _abbrev_one(raw)
            primary = _is_primary(raw.replace('{', '').replace('}', ''))
            if i > 0:
                _run(para, ", ")
            _run(para, abbr, underline=primary)

    # Year
    year_str = (e.get('year') or '').strip()
    _run(para, f" ({year_str}), " if year_str else ", ")

    # Title in blue with typographic quotes
    title = _clean_title(apply_title_case(_prep_title(e.get('title') or 'Untitled')))
    _run(para, "\u201c")
    _run(para, title, color=C_TITLE)
    _run(para, "\u201d")

    # Note
    note = note_clean(e.get('note') or '')
    if note:
        _run(para, f" [{note}]", color=C_NOTE)

    # Venue (italic) — inline, same sentence
    venue_full  = _shorten_venue(pub_venue_str(e))
    venue_clean = re.sub(r'\s*\(\d{4}\)\s*$', '', venue_full).strip()
    if venue_clean:
        _run(para, ", ")
        _run(para, venue_clean + ".", italic=True)
    else:
        _run(para, ".")

# ---------------------------------------------------------------------------
# Footer with page numbers
# ---------------------------------------------------------------------------

def _add_footer(doc: Document):
    footer = doc.sections[0].footer
    para   = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Left: author name
    r1 = para.add_run(f"{AUTHOR_NAME}  ·  Publication List    ")
    r1.font.name = F_BODY; r1.font.size = Pt(8); r1.font.color.rgb = C_LIGHT; r1.font.italic = True
    # Page number field
    run = para.add_run()
    run.font.name = F_BODY; run.font.size = Pt(8); run.font.color.rgb = C_LIGHT
    for tag, text in [('begin', None), ('instrText', ' PAGE '), ('end', None)]:
        if tag == 'instrText':
            el = OxmlElement('w:instrText'); el.text = text
        else:
            el = OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'), tag)
        run._r.append(el)

# ---------------------------------------------------------------------------
# Document build
# ---------------------------------------------------------------------------

def build_docx(entries: list, output_path: str):
    doc = Document()

    # Page setup: A4
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(1.8)
    sec.right_margin  = Cm(1.8)
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    # Clear default empty paragraph
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    # Masthead
    h = doc.add_paragraph()
    r = h.add_run("Publication List")
    r.font.name = F_BODY; r.font.size = Pt(22); r.font.color.rgb = C_HDR; r.font.bold = False

    sub = doc.add_paragraph()
    r = sub.add_run(f"{len(entries)} publications  ·  {AUTHOR_NAME}")
    r.font.name = F_BODY; r.font.size = Pt(10.5); r.font.italic = True; r.font.color.rgb = C_LIGHT
    sub.paragraph_format.space_after = Pt(12)

    _add_footer(doc)

    # Sort entries by year (desc), then conf → journal → book within year
    years_dict: OrderedDict = OrderedDict()
    for e in entries:
        yr = (e.get('year') or 'Unknown').strip()
        years_dict.setdefault(yr, []).append(e)
    for yr in years_dict:
        years_dict[yr].sort(key=_entry_kind)

    ordered = [(yr, e) for yr, yr_entries in years_dict.items() for e in yr_entries]
    total   = len(ordered)
    for i, (yr, e) in enumerate(ordered):
        add_entry(doc, e, total - i)

    # Timestamp
    ts = doc.add_paragraph()
    ts.paragraph_format.space_before = Pt(18)
    r = ts.add_run(datetime.now().strftime("Generated %d %B %Y, %H:%M"))
    r.font.name = F_BODY; r.font.size = Pt(8); r.font.italic = True; r.font.color.rgb = C_LIGHT

    doc.save(output_path)

# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print("Parsing bib file…")
    entries = parse_bib()
    print(f"  {len(entries)} entries found")
    print(f"  Writing {OUTPUT_DOCX}…")
    build_docx(entries, OUTPUT_DOCX)
    print(f"\nSaved to: {OUTPUT_DOCX}")

if __name__ == "__main__":
    main()
